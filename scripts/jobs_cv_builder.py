#!/usr/bin/env python3
"""Genera CV ATS en Word/PDF adaptado a una vacante pegada o archivo JD."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from jobs_common import JOBS_WS, ROOT, load_config

CV_PROFILE_PATH = ROOT / "config/jobs/cv_profile.json"
BRAND = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)

UI_NOISE = re.compile(
    r"(?i)^(mis postulaciones|mi perfil|cerrar sesi[oó]n|configuraci[oó]n|"
    r"informaci[oó]n personal|video presentaci[oó]n|powered by|help_outline|"
    r"list|person|settings|logout|sobre aira|t[eé]rminos y condiciones|"
    r"preguntas frecuentes|©.*aira.*)$"
)

ROLE_PATTERNS: list[tuple[str, str, list[str]]] = [
    ("data_analyst", "Data Analyst", ["data analyst", "analista de datos", "people analytics", "hr analytics", "analista hr"]),
    ("data_engineer", "Data Engineer", ["data engineer", "ingeniería de datos", "ingeniero de datos", "etl", "big data"]),
    ("devsecops", "Senior DevSecOps Engineer", ["devsecops", "ciberseguridad", "cybersecurity", "pentesting"]),
    ("sre", "Senior SRE / Platform Engineer", ["site reliability", "sre", "platform engineer", "reliability engineer"]),
    ("technical_lead", "Senior Technical Lead", ["technical lead", "tech lead", "líder técnico", "arquitecto tecnológico", "arquitecto tecnologico"]),
    ("mlops", "Senior MLOps / AI Engineer", ["mlops", "ai engineer", "machine learning engineer", "llm", "ai platform"]),
    ("devops", "Senior DevOps / Cloud Engineer", ["devops", "cloud engineer", "cloud architect", "platform engineer"]),
]

HEADLINES: dict[str, str] = {
    "data_analyst": "Data Analyst | SQL | Power BI | Python | HR Analytics & People Data",
    "data_engineer": "Data Engineer | SQL | Python | ETL | Cloud Data Pipelines",
    "devsecops": "Senior DevSecOps Engineer | Cloud Security | CI/CD | AWS",
    "sre": "Senior SRE / Platform Engineer | Observability | Kubernetes | AWS",
    "technical_lead": "Technical Lead | Cloud Architecture | APIs | Integration",
    "mlops": "Senior MLOps / AI Engineer | Python | AWS | LLM & ML Pipelines",
    "devops": "Senior DevOps / Cloud Engineer | AWS | Kubernetes | Terraform | CI/CD",
}


@dataclass
class VacancyInfo:
    raw: str
    title: str = "Oferta laboral"
    company: str = ""
    location: str = ""
    description: str = ""
    requirements: str = ""
    keywords: list[str] = field(default_factory=list)
    role_key: str = "devops"
    role_label: str = "Senior DevOps / Cloud Engineer"

    @property
    def slug(self) -> str:
        base = f"{self.company}-{self.title}-{self.role_key}"
        digest = hashlib.sha1(self.raw.encode()).hexdigest()[:8]
        slug = re.sub(r"[^a-z0-9]+", "_", base.lower()).strip("_")[:60]
        return f"{slug}_{digest}" if slug else digest


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def load_cv_profile() -> dict[str, Any]:
    if CV_PROFILE_PATH.exists():
        return json.loads(CV_PROFILE_PATH.read_text(encoding="utf-8"))
    return {}


def strip_ui_noise(text: str) -> str:
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if UI_NOISE.match(stripped):
            continue
        if stripped in {"MC", "Descripción", "Oferta", "Requisitos", "Candidato/a", "Proceso de", "Selección", "Sobre", "Nosotros"}:
            continue
        lines.append(stripped)
    return "\n".join(lines)


def extract_section(text: str, start_labels: tuple[str, ...], end_labels: tuple[str, ...]) -> str:
    low = text.lower()
    start_idx = -1
    for label in start_labels:
        idx = low.find(label.lower())
        if idx >= 0:
            start_idx = idx + len(label)
            break
    if start_idx < 0:
        return ""
    end_idx = len(text)
    for label in end_labels:
        idx = low.find(label.lower(), start_idx)
        if idx >= 0:
            end_idx = min(end_idx, idx)
    return clean_text(text[start_idx:end_idx])


def infer_company(lines: list[str], text: str) -> str:
    for i, line in enumerate(lines[:12]):
        if " - " in line and i > 0:
            left = line.split(" - ", 1)[0].strip()
            if 2 < len(left) < 80 and not re.search(r"(?i)full time|presencial|remoto|publicado", left):
                return left
        if i > 0 and re.match(r"^[A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÑáéíóúñ0-9 .&()/-]{2,60}$", line):
            if not re.search(r"(?i)data analyst|devops|engineer|analyst|developer|architect|lead|sre", line):
                if lines[i - 1] == line or (i > 1 and lines[i - 2] == line):
                    return line
    match = re.search(r"(?i)(?:empresa|company)\s*[:\-]\s*([^\n]{2,80})", text)
    if match:
        return clean_text(match.group(1))[:80]
    return ""


def infer_title(lines: list[str], text: str) -> str:
    role_terms = (
        "data analyst", "devops", "sre", "cloud engineer", "software engineer",
        "technical lead", "architect", "data engineer", "mlops", "ai engineer",
        "platform engineer", "analyst", "developer", "ingeniero",
    )
    for line in lines[:20]:
        low = line.lower()
        if any(term in low for term in role_terms) and len(line) < 90:
            if not re.search(r"(?i)full time|presencial|remoto|publicado|permanente", line):
                return line.strip()
    match = re.search(
        r"(?i)(data analyst|devops engineer|cloud engineer|sre|technical lead|"
        r"data engineer|mlops engineer|software engineer|platform engineer|"
        r"arquitecto tecnol[oó]gico)[^\n]{0,40}",
        text,
    )
    if match:
        return clean_text(match.group(0))[:90]
    return "Oferta laboral"


def infer_location(text: str) -> str:
    match = re.search(r"🇨🇱\s*([^|\n]{3,60})", text)
    if match:
        return clean_text(match.group(1))
    match = re.search(r"(?i)(santiago|las condes|providencia|metropolitana|chile)[^.\n]{0,40}", text)
    if match:
        return clean_text(match.group(0))[:60]
    return ""


def infer_role_key(text: str) -> tuple[str, str]:
    low = text.lower()
    best_key = "devops"
    best_label = HEADLINES["devops"].split("|")[0].strip()
    best_hits = 0
    for key, label, terms in ROLE_PATTERNS:
        hits = sum(1 for term in terms if term in low)
        if hits > best_hits:
            best_hits = hits
            best_key = key
            best_label = label
    if "hr" in low and ("analyst" in low or "datos" in low or "people" in low):
        best_key = "data_analyst"
        best_label = "Data Analyst"
    return best_key, best_label


def extract_keywords(text: str, limit: int = 24) -> list[str]:
    low = text.lower()
    catalog = [
        "sql", "power bi", "python", "pandas", "kpi", "dashboard", "etl", "hr analytics",
        "people analytics", "power query", "postgresql", "excel", "estadística", "reporting",
        "aws", "kubernetes", "terraform", "docker", "jenkins", "ci/cd", "devops", "sre",
        "observability", "prometheus", "grafana", "lambda", "microservices", "apis",
        "confidencialidad", "calidad de datos", "insights", "integración", "automatización",
    ]
    found = [kw for kw in catalog if kw in low]
    tokens = re.findall(r"[a-záéíóúñ0-9+#./-]{4,}", low)
    for token in tokens:
        if token not in found and len(found) < limit:
            if any(x in token for x in ("data", "cloud", "anal", "report", "dash", "kpi")):
                found.append(token)
    return found[:limit]


def parse_vacancy(text: str) -> VacancyInfo:
    cleaned = strip_ui_noise(text)
    lines = [l for l in cleaned.splitlines() if l.strip()]
    body = clean_text(cleaned)
    role_key, role_label = infer_role_key(body)
    info = VacancyInfo(
        raw=body,
        title=infer_title(lines, body),
        company=infer_company(lines, body),
        location=infer_location(body),
        description=extract_section(body, ("descripción", "oferta", "description"), ("requisitos", "requirements", "proceso de", "sobre nosotros")),
        requirements=extract_section(body, ("requisitos", "candidato", "requirements"), ("proceso de", "sobre nosotros", "sobre aira", "©")),
        keywords=extract_keywords(body),
        role_key=role_key,
        role_label=role_label,
    )
    if not info.description:
        info.description = body[:2500]
    return info


def role_skill_key(role_key: str) -> str:
    if role_key in {"data_analyst"}:
        return "data_analyst"
    if role_key in {"data_engineer"}:
        return "data_engineer"
    if role_key in {"technical_lead", "mlops"}:
        return "technical_lead"
    return "devops"


def focus_phrase(vacancy: VacancyInfo) -> str:
    kws = vacancy.keywords[:8]
    if kws:
        return ", ".join(kws)
    return HEADLINES.get(vacancy.role_key, HEADLINES["devops"]).split("|")[1:4].__str__()


def de_emphasize_note(role_key: str) -> str:
    if role_key == "data_analyst":
        return ' Mantendré tu seniority, pero sin que el CV "grite DevOps" como perfil principal.'
    if role_key in {"devops", "sre", "devsecops"}:
        return " Mantendré tu seniority cloud/DevOps como eje principal."
    return ""


def whatsapp_intro(vacancy: VacancyInfo) -> str:
    company = vacancy.company or "la empresa"
    focus = focus_phrase(vacancy)
    return (
        f"Voy a generar una versión ATS enfocada 100% en {vacancy.role_label} para {company}: "
        f"{focus}.{de_emphasize_note(vacancy.role_key)}"
    )


def whatsapp_done(vacancy: VacancyInfo, docx_path: Path, pdf_path: Path | None) -> str:
    company = vacancy.company or "target"
    lines = [
        whatsapp_intro(vacancy),
        "",
        f"Listo. Generé una versión del CV ajustada a la vacante {vacancy.title} - {company}, "
        f"enfocada en {focus_phrase(vacancy)}.",
        "",
        "Descargas:",
        f"CV en Word (.docx): {docx_path}",
    ]
    if pdf_path and pdf_path.exists():
        lines.append(f"CV en PDF: {pdf_path}")
    else:
        lines.append("CV en PDF: no disponible (instalar libreoffice-writer-nogui para exportar PDF).")
    return "\n".join(lines)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "000000", size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def set_borders(table, color: str = "D9E2F3", sz: str = "4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), sz)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)
        borders.append(elem)
    tbl_pr.append(borders)


def add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BRAND
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullets(doc: Document, items: list[str], size: float = 8.5) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(0.7)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(item)
        r.font.size = Pt(size)


def add_role(doc: Document, role: str, company: str, dates: str, summary: str, bullets: list[str], tech: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{company} - {role}")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = BRAND
    r2 = p.add_run(f" | {dates}")
    r2.font.size = Pt(8.5)
    r2.italic = True
    if summary:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.line_spacing = 1.0
        rr = p2.add_run(summary)
        rr.font.size = Pt(8.7)
    add_bullets(doc, bullets, size=8.5)
    if tech:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(1)
        p3.paragraph_format.line_spacing = 1.0
        rt = p3.add_run("Tecnologías: ")
        rt.bold = True
        rt.font.size = Pt(8.2)
        rt2 = p3.add_run(tech)
        rt2.font.size = Pt(8.2)
        rt2.font.color.rgb = MUTED


def alignment_table(vacancy: VacancyInfo) -> tuple[list[str], list[str]]:
    if vacancy.role_key == "data_analyst":
        return (
            ["SQL / Datos", "Power BI / KPIs", "Python / Automatización", "HR Analytics / Confidencialidad"],
            [
                "Extracción, limpieza, integración, modelamiento relacional y optimización de consultas.",
                "Dashboards ejecutivos, reportes de gestión, indicadores operacionales y visualización.",
                "Procesamiento de datos, validaciones, APIs, automatización de reportes y flujos repetibles.",
                "KPIs de personas, insights accionables, calidad del dato y manejo responsable de información sensible.",
            ],
        )
    return (
        ["Cloud / DevOps", "CI/CD / IaC", "Observabilidad", "Seguridad / FinOps"],
        [
            "AWS, contenedores, plataforma y operación de sistemas críticos.",
            "Terraform, Jenkins, GitHub Actions y despliegues automatizados.",
            "Métricas DORA, monitoreo, trazabilidad y respuesta a incidentes.",
            "DevSecOps, controles de acceso, entornos regulados y eficiencia de costos.",
        ],
    )


def build_docx(vacancy: VacancyInfo, out_path: Path, profile: dict[str, Any] | None = None) -> Path:
    profile = profile or load_cv_profile()
    skill_key = role_skill_key(vacancy.role_key)
    experience_key = skill_key if skill_key != "technical_lead" else vacancy.role_key
    if experience_key not in {"data_analyst", "devops", "technical_lead", "data_engineer"}:
        experience_key = "devops"

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.35)
    sec.bottom_margin = Inches(0.35)
    sec.left_margin = Inches(0.48)
    sec.right_margin = Inches(0.48)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(8.8)
    styles["List Bullet"].font.name = "Arial"
    styles["List Bullet"].font.size = Pt(8.6)

    headline = HEADLINES.get(vacancy.role_key, HEADLINES["devops"])
    owner = profile.get("owner") or load_config().get("owner") or "Mauricio Castro"
    contact = (
        f"{profile.get('location', 'Santiago, Chile')} | {profile.get('phone', '')} | "
        f"{profile.get('email', '')} | {profile.get('linkedin', '')} | "
        f"{profile.get('github', '')} | {profile.get('website', '')}"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(owner)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(headline)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(64, 64, 64)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(contact)
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED

    add_section_heading(doc, "Perfil profesional")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(profile.get("summary_base", ""))
    r.font.size = Pt(8.9)

    company = vacancy.company or "target"
    add_section_heading(doc, f"Resumen ejecutivo para {company} - {vacancy.title}")
    headers, vals = alignment_table(vacancy)
    table = doc.add_table(rows=2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    for i, header in enumerate(headers):
        shade_cell(table.cell(0, i), "1F4E79")
        set_cell_text(table.cell(0, i), header, bold=True, color="FFFFFF", size=7.7)
        table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for i, val in enumerate(vals):
        shade_cell(table.cell(1, i), "F3F6FA")
        set_cell_text(table.cell(1, i), val, size=7.5)
        table.cell(1, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    add_section_heading(doc, "Competencias técnicas")
    for title, desc in profile.get("skill_blocks", {}).get(skill_key, []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0.6)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f"{title}: ")
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = BRAND
        r2 = p.add_run(desc)
        r2.font.size = Pt(8.5)

    add_section_heading(doc, "Experiencia profesional relevante")
    for exp in profile.get("experiences", []):
        role_data = (exp.get("roles") or {}).get(experience_key) or (exp.get("roles") or {}).get("devops")
        if not role_data:
            continue
        add_role(
            doc,
            role_data.get("title", ""),
            exp.get("company", ""),
            role_data.get("dates", ""),
            role_data.get("summary", ""),
            role_data.get("bullets", []),
            role_data.get("tech", ""),
        )

    value_title = "Logros y aporte para HR Analytics" if vacancy.role_key == "data_analyst" else "Logros y aporte clave"
    add_section_heading(doc, value_title)
    add_bullets(doc, profile.get("value_props", {}).get(skill_key, profile.get("value_props", {}).get("devops", [])), size=8.5)

    add_section_heading(doc, "Educación y certificaciones")
    add_bullets(doc, profile.get("education", []), size=8.5)

    add_section_heading(doc, "Idiomas")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(profile.get("languages", ""))
    r.font.size = Pt(8.6)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        f"CV optimizado para {vacancy.title} - {company} | {focus_phrase(vacancy)}"
    )
    fr.font.size = Pt(7)
    fr.font.color.rgb = RGBColor(127, 127, 127)

    props = doc.core_properties
    props.author = owner
    props.title = f"CV {owner} - {vacancy.title} {company}"
    props.subject = f"CV ATS {vacancy.role_label}"
    props.keywords = ", ".join(vacancy.keywords)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def _libreoffice_bin() -> str | None:
    for candidate in ("libreoffice", "soffice"):
        path = shutil.which(candidate)
        if path and "orca" not in path:
            return path
    return None


def docx_to_pdf(docx_path: Path, pdf_path: Path | None = None) -> Path | None:
    pdf_path = pdf_path or docx_path.with_suffix(".pdf")
    lo = _libreoffice_bin()
    if lo:
        proc = subprocess.run(
            [lo, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        generated = docx_path.with_suffix(".pdf")
        if proc.returncode == 0 and generated.exists():
            if generated != pdf_path:
                generated.replace(pdf_path)
            return pdf_path
    return _pdf_fallback(docx_path, pdf_path)


def _extract_docx_text(path: Path) -> str:
    import zipfile
    from xml.etree import ElementTree

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    parts = [node.text for node in root.findall(".//w:t", ns) if node.text]
    return clean_text(" ".join(parts))


def _pdf_fallback(docx_path: Path, pdf_path: Path) -> Path | None:
    try:
        from fpdf import FPDF
    except ImportError:
        return None

    text = _extract_docx_text(docx_path)
    if not text:
        return None

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font("Helvetica", size=9)
    for paragraph in re.split(r"\s{2,}|\n+", text):
        chunk = clean_text(paragraph)
        if not chunk:
            continue
        pdf.multi_cell(0, 4.5, chunk)
        pdf.ln(1)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(pdf_path))
    return pdf_path if pdf_path.exists() else None


def output_paths(vacancy: VacancyInfo, out_dir: Path | None = None) -> tuple[Path, Path]:
    out_dir = out_dir or (JOBS_WS / "cv_generated" / vacancy.slug)
    out_dir.mkdir(parents=True, exist_ok=True)
    company_slug = re.sub(r"[^a-zA-Z0-9]+", "_", (vacancy.company or "Target").strip("_"))[:30]
    role_slug = re.sub(r"[^a-zA-Z0-9]+", "_", vacancy.role_label.strip("_"))[:25]
    base = f"CV_Mauricio_Castro_{role_slug}_{company_slug}"
    return out_dir / f"{base}.docx", out_dir / f"{base}.pdf"


def generate_cv_files(text: str, out_dir: Path | None = None) -> dict[str, Any]:
    vacancy = parse_vacancy(text)
    docx_path, pdf_path = output_paths(vacancy, out_dir)
    build_docx(vacancy, docx_path)
    pdf_result = docx_to_pdf(docx_path, pdf_path)
    return {
        "vacancy": {
            "title": vacancy.title,
            "company": vacancy.company,
            "location": vacancy.location,
            "role_key": vacancy.role_key,
            "role_label": vacancy.role_label,
            "keywords": vacancy.keywords,
            "slug": vacancy.slug,
        },
        "docx": str(docx_path),
        "pdf": str(pdf_result) if pdf_result else "",
        "whatsapp_intro": whatsapp_intro(vacancy),
        "whatsapp_reply": whatsapp_done(vacancy, docx_path, pdf_result),
    }
